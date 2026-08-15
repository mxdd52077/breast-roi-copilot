from __future__ import annotations

from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "artifacts"
ASSET_DIR = OUT_DIR / "prd_assets"
OUT_DIR.mkdir(exist_ok=True)
ASSET_DIR.mkdir(exist_ok=True)
OUTPUT = OUT_DIR / "APEX乳腺筛查循证智能决策平台_PRD_V1.0.docx"

NAVY = "17233C"
TEAL = "087F83"
LIGHT_TEAL = "E8F4F3"
BLUE = "2C6E9B"
LIGHT_BLUE = "EAF2F8"
GRAY = "667085"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "D7DEE7"
RED = "B42318"
LIGHT_RED = "FDECEC"
GOLD = "A66A00"
LIGHT_GOLD = "FFF5DB"
WHITE = "FFFFFF"


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeatable_header(section, text: str) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), MID_GRAY)
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("— ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    paragraph.add_run(" —")
    for r in paragraph.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(GRAY)


def set_col_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None,
              header_color: str = TEAL, font_size: float = 8.7) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, h in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_fill(cell, header_color)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        r.font.size = Pt(font_size)
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if len(table.rows) % 2 == 1:
                set_cell_fill(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            r.font.size = Pt(font_size)
    if widths:
        set_col_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def add_callout(doc: Document, title: str, body: str, kind: str = "info") -> None:
    colors = {
        "info": (LIGHT_BLUE, BLUE),
        "success": (LIGHT_TEAL, TEAL),
        "warning": (LIGHT_GOLD, GOLD),
        "risk": (LIGHT_RED, RED),
    }
    fill, accent = colors[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16.5)
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(accent)
    r2 = p.add_run(body)
    r2.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for label, value in rows:
        row = table.add_row()
        prevent_row_split(row)
        set_cell_fill(row.cells[0], LIGHT_TEAL)
        for cell in row.cells:
            set_cell_margins(cell, top=110, bottom=110)
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.color.rgb = RGBColor.from_string(TEAL)
        row.cells[1].paragraphs[0].add_run(value)
    set_col_widths(table, [4.2, 12.3])
    doc.add_paragraph()


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "/System/Library/Fonts/STHeiti Medium.ttc" if bold else "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def _centered(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill: str, spacing: int = 6) -> None:
    bounds = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
    width = bounds[2] - bounds[0]
    height = bounds[3] - bounds[1]
    x = box[0] + (box[2] - box[0] - width) / 2
    y = box[1] + (box[3] - box[1] - height) / 2
    draw.multiline_text((x, y), text, font=font, fill=fill, spacing=spacing, align="center")


def create_architecture(path: Path) -> None:
    image = Image.new("RGB", (2160, 1040), "white")
    draw = ImageDraw.Draw(image)
    _centered(draw, (0, 25, 2160, 115), "APEX 乳腺筛查循证智能决策平台｜功能架构", _font(43, True), "#17233C")
    layers = [
        (145, "用户入口", ["AI决策助手", "数据接入", "证据与参数", "决策分析"]),
        (365, "产品编排层", ["场景解析", "缺失项路由", "人工确认", "报告编排"]),
        (585, "能力层", ["PubMed检索\nAI提取", "确定性ROI引擎", "外展优先级规则", "本地校验器"]),
        (805, "数据与治理", ["医院CSV\n（会话级）", "证据价值库", "参数来源记录", "审计与安全边界"]),
    ]
    for y, label, items in layers:
        draw.text((38, y + 50), label, font=_font(27, True), fill="#17233C")
        for i, item in enumerate(items):
            x = 390 + i * 430
            box = (x, y, x + 365, y + 145)
            fill = "#EAF2F8" if y == 585 else "#E8F4F3"
            outline = "#2C6E9B" if y == 585 else "#087F83"
            draw.rounded_rectangle(box, radius=22, fill=fill, outline=outline, width=3)
            _centered(draw, box, item, _font(26, True), "#17233C")
    for y in (290, 510, 730):
        draw.line((1080, y, 1080, y + 55), fill="#98A2B3", width=5)
        draw.polygon([(1068, y + 50), (1092, y + 50), (1080, y + 70)], fill="#98A2B3")
    image.save(path)


def create_flow(path: Path) -> None:
    image = Image.new("RGB", (2160, 850), "white")
    draw = ImageDraw.Draw(image)
    _centered(draw, (0, 25, 2160, 110), "核心闭环：医院数据 → 补充证据 → 人工确认 → 确定性计算 → AI报告", _font(39, True), "#17233C")
    steps = [
        ("1", "上传医院CSV", "校验、去标识化\n生成派生字段"),
        ("2", "检查缺失参数", "医院数据优先\n只补缺失项"),
        ("3", "检索与提取证据", "PubMed/官方来源\nAI仅生成候选"),
        ("4", "人工确认参数", "保留来源与假设\n不自动写入模型"),
        ("5", "运行ROI与资源优化", "原R公式迁移\n透明排序规则"),
        ("6", "生成管理层报告", "AI撰写 + 本地校验\n不合格则拦截"),
    ]
    for i, (num, title, desc) in enumerate(steps):
        x = 45 + i * 350
        box = (x, 180, x + 290, 575)
        draw.rounded_rectangle(box, radius=22, fill="#F7FAFC", outline="#087F83", width=3)
        draw.ellipse((x + 115, 205, x + 175, 265), fill="#087F83")
        _centered(draw, (x + 115, 205, x + 175, 265), num, _font(24, True), "white")
        _centered(draw, (x + 15, 285, x + 275, 360), title, _font(25, True), "#17233C")
        _centered(draw, (x + 20, 375, x + 270, 535), desc, _font(22), "#667085", spacing=12)
        if i < len(steps) - 1:
            draw.line((x + 295, 380, x + 337, 380), fill="#98A2B3", width=5)
            draw.polygon([(x + 330, 368), (x + 350, 380), (x + 330, 392)], fill="#98A2B3")
    _centered(draw, (80, 650, 2080, 785), "关键边界：LLM 不执行 ROI 数学计算，不自动修改参数，\n不显示未通过数字与引用校验的报告。", _font(27, True), "#B42318", spacing=10)
    image.save(path)


def configure_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(1.9)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.15)
    sec.right_margin = Cm(2.15)
    set_repeatable_header(sec, "APEX · 乳腺筛查循证智能决策平台｜PRD V1.0")
    add_page_number(sec.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Hiragino Sans GB"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 12),
        ("Subtitle", 13, TEAL, 0, 8),
        ("Heading 1", 19, NAVY, 14, 8),
        ("Heading 2", 14, TEAL, 10, 5),
        ("Heading 3", 11.5, BLUE, 7, 4),
    ):
        st = doc.styles[style_name]
        st.font.name = "Hiragino Sans GB"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        st = doc.styles[style_name]
        st.font.name = "Hiragino Sans GB"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
        st.font.size = Pt(10)


def apply_cjk_font(doc: Document) -> None:
    """Force a LibreOffice-compatible CJK font on every text run."""
    def style_run(run) -> None:
        run.font.name = "Hiragino Sans GB"
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), "Hiragino Sans GB")

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            style_run(run)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                for run in paragraph.runs:
                    style_run(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        style_run(run)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(62)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("APEX")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("乳腺筛查循证智能决策平台")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("产品需求文档（PRD）")
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run("医院数据优先 · 循证参数补充 · 确定性 ROI · AI 管理层报告")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(GRAY)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    values = [("文档版本", "V1.0"), ("评审状态", "初稿"), ("撰写日期", "2026-08-15"), ("保密级别", "作品集 / 内部评审")]
    for row, (k, v) in zip(table.rows, values):
        set_cell_fill(row.cells[0], LIGHT_TEAL)
        set_cell_margins(row.cells[0], top=120, bottom=120)
        set_cell_margins(row.cells[1], top=120, bottom=120)
        a = row.cells[0].paragraphs[0].add_run(k)
        a.bold = True
        a.font.color.rgb = RGBColor.from_string(TEAL)
        row.cells[1].paragraphs[0].add_run(v)
    set_col_widths(table, [4.5, 8.5])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    r = p.add_run("基于当前 Streamlit MVP 与原 R Shiny 乳腺筛查 ROI 模型整理")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_page_break()


def build_doc() -> None:
    arch = ASSET_DIR / "architecture.png"
    flow = ASSET_DIR / "core_flow.png"
    create_architecture(arch)
    create_flow(flow)

    doc = Document()
    configure_doc(doc)
    add_cover(doc)

    doc.add_heading("阅读说明", level=1)
    add_callout(doc, "文档口径", "本 PRD 以当前精简版 MVP 为准。已实现能力以“当前”描述；尚未实现但建议补充的能力标记为 P1/P2 或“待确认”。所有真实业务数据、生产并发与运营效果均不做推测。", "info")
    doc.add_paragraph("本期可见主线仅保留四个模块：AI 决策助手、数据接入、证据与参数、决策分析。历史独立页面（管理层报告、模型性能评估、循证问答、评测与 Bad Case）已从普通用户导航隐藏，但底层代码与测试资产保留，供专业复核和后续迭代使用。")
    doc.add_heading("目录", level=2)
    for item in [
        "一、文档信息", "二、需求背景与目标", "三、需求范围与功能清单", "四、需求内容详述",
        "五、产品功能详细说明", "六、原型与交互说明", "七、非功能需求（NFR）",
        "八、数据埋点与分析", "九、上线与运营", "十、风险与兜底方案", "十一、附录",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    doc.add_heading("一、文档信息", level=1)
    doc.add_heading("1.1 基础信息", level=2)
    add_kv_table(doc, [
        ("文档名称", "APEX 乳腺筛查循证智能决策平台产品需求文档"),
        ("文档版本", "V1.0"),
        ("撰写人", ""),
        ("所属业务线 / 产品", "医疗健康 / AI 决策支持 / APEX"),
        ("撰写日期", "2026-08-15"),
        ("更新日期", "2026-08-15"),
        ("评审状态", "初稿"),
    ])
    doc.add_heading("1.2 变更历史", level=2)
    add_table(doc, ["版本", "日期", "修改人", "修改内容"], [["V1.0", "2026-08-15", "", "基于精简版 MVP 创建完整 PRD 初稿"]], [2, 3, 3, 8.5])

    doc.add_heading("二、需求背景与目标", level=1)
    doc.add_heading("2.1 业务背景", level=2)
    doc.add_paragraph("乳腺筛查项目的扩容决策通常同时涉及目标人群、当前覆盖率、筛查方式、检出率、召回率、随访完成率、分期治疗成本等多类参数。医院运营或项目负责人需要从院内报表、PubMed 文献、SEER/CMS/指南等来源中查找信息，再进行手工测算，流程长、参数来源难追溯，且非专业使用者难以解释模型结果。")
    doc.add_paragraph("现有 R Shiny 工具已经具备确定性的乳腺筛查 ROI 数值模型，但原形态偏分析工具：参数较多、证据补充与计算割裂、报告依赖人工整理。当前版本将其产品化为“医院数据优先”的 AI 决策工作流：系统先读取院内数据，只针对缺失参数引导检索外部依据，经人工确认后调用原 R 公式迁移的 Python 引擎计算 ROI，并由大模型生成经本地校验的管理层报告。")
    doc.add_heading("2.2 核心问题 / 痛点", level=2)
    add_bullet(doc, "用户痛点：普通运营或互联网背景的评审者难以理解大量医学与生物统计参数，不知道先上传什么、缺什么、到哪里补充，也难以把计算结果转化为管理层可读结论。")
    add_bullet(doc, "业务痛点：有限外展资源无法快速识别优先联系对象；ROI 结果与证据、参数来源、人工决策之间缺乏统一链路。")
    add_bullet(doc, "数据痛点：医院原始字段与下游模型字段不一致；患者级数据存在缺失、格式、日期和隐私风险；文献证据与本地运营数据适用范围不同。")
    add_bullet(doc, "AI 风险：大模型可能生成不存在的 PMID、篡改 ROI 数字、将情景结果误写成真实效果，或给出缺乏证据的确定性结论。")
    doc.add_heading("2.3 产品目标（可量化）", level=2)
    add_table(doc, ["目标", "指标口径", "V1.0 验收目标", "当前状态"], [
        ["完成一站式决策闭环", "标准演示场景从数据应用到报告草稿的完成率", "≥95%（内部验收集）", "待建立正式验收记录"],
        ["保证计算一致性", "核心 ROI 回归用例通过率", "100%", "现有项目自动化测试 75/75 通过"],
        ["控制 AI 幻觉", "报告展示前数字与 PMID/摘录校验通过率", "100%；未通过不得展示", "已实现本地阻断校验"],
        ["降低使用门槛", "首次使用者完成主流程所需页面数", "≤4 个可见模块", "已完成导航精简"],
        ["保障人工决策权", "AI 参数自动写入 ROI 的次数", "0", "已实现人工确认门槛"],
    ], [3.2, 6.2, 4.2, 3.4], font_size=8.2)
    add_callout(doc, "指标说明", "以上为产品验收目标，不代表真实医院上线效果；生产转化率、节省金额和临床获益必须在真实部署后另行验证。", "warning")

    doc.add_heading("三、需求范围与功能清单", level=1)
    doc.add_heading("3.1 产品定位与目标用户", level=2)
    add_table(doc, ["角色", "核心任务", "主要使用模块", "权限边界"], [
        ["医院运营 / 项目负责人（主用户）", "建立筛查扩容场景、确认参数、查看 ROI 与报告", "AI 决策助手、数据接入", "可应用数据、确认参数、生成草稿"],
        ["数据分析师", "检查数据质量、复核计算与外展排序", "数据接入、高级 ROI 仿真、外展资源优化", "可查看派生字段和计算过程"],
        ["循证研究 / 医学顾问", "检索文献、审核效应量和适用范围", "证据与参数", "批准或拒绝证据候选"],
        ["管理层 / 面试评审者", "快速理解场景价值、证据边界和下一步行动", "AI 决策助手生成的报告草稿", "只读为主"],
    ], [3.4, 5.2, 4.1, 4.2], font_size=8.2)
    doc.add_heading("3.2 本期范围", level=2)
    add_bullet(doc, "仅支持 40–74 岁女性的 DBT / 3D mammography 筛查场景。")
    add_bullet(doc, "支持医院 CSV 会话级上传、质量校验、系统派生字段和非合成数据患者 ID 本地去标识化。")
    add_bullet(doc, "支持医院数据优先的参数缺失检查、定向来源指引、人工确认、确定性 ROI 计算与 AI 管理层报告。")
    add_bullet(doc, "支持真实 PubMed 检索、AI 效应量候选提取、人工审核和证据价值库存储。")
    add_bullet(doc, "支持在固定外展名额下对比随机联系与规则优先联系的预期价值。")
    doc.add_heading("3.3 非本期范围", level=2)
    add_bullet(doc, "MRI、超声、数字乳腺摄影等其他筛查方式的独立参数模型。")
    add_bullet(doc, "临床诊断、个体化治疗建议、患者风险诊断或自动外展执行。")
    add_bullet(doc, "生产级 EHR/HDR/FHIR 集成、真实患者数据持久化、多租户权限与医院级单点登录。")
    add_bullet(doc, "让大模型直接计算 ROI、自动批准证据、自动修改模型参数或承诺临床/财务结果。")
    add_bullet(doc, "多癌种扩展与真实世界因果效果验证。")
    doc.add_heading("3.4 整体功能架构", level=2)
    doc.add_picture(str(arch), width=Inches(6.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("图 1｜产品功能架构。用户看到的是简化工作流；检索、校验、计算和审计能力在后台解耦。", style="Caption")
    doc.add_heading("3.5 功能清单与优先级", level=2)
    add_table(doc, ["模块", "功能点", "优先级", "实现状态 / 备注"], [
        ["AI 决策助手", "读取已应用医院数据并建立场景", "P0", "已实现；无数据时阻断并引导上传"],
        ["AI 决策助手", "识别已知/缺失参数并给出来源路线", "P0", "已实现；医院数据优先"],
        ["AI 决策助手", "人工确认后调用原 R 公式迁移引擎", "P0", "已实现；LLM 不参与数学计算"],
        ["AI 决策助手", "生成并校验管理层报告", "P0", "已实现；支持实时 AI 与验证演示"],
        ["数据接入", "CSV 上传、字段/日期/年龄/重复值校验", "P0", "已实现；单文件≤100,000行"],
        ["数据接入", "派生 Care Gap、完成概率、检出概率", "P0", "已实现；规则/假设来源可见"],
        ["证据与参数", "PubMed 检索与文献查看", "P0", "已实现；NCBI E-utilities"],
        ["证据与参数", "AI 效应量提取 + 人工审核 + 价值库", "P1", "已实现专业入口；普通用户可跳过"],
        ["高级 ROI 仿真", "核心三参数 what-if 与结果卡", "P1", "已实现；专业复核工具"],
        ["外展资源优化", "随机 vs 优先联系策略对比", "P1", "已实现；规则排序，不是 LLM"],
        ["运营分析", "埋点、漏斗、AI 成本与延迟看板", "P1", "待开发"],
        ["模型扩展", "数字乳腺摄影独立参数配置", "P2", "待独立证据与回归验证后开发"],
    ], [3.0, 6.2, 1.7, 5.1], font_size=7.9)

    doc.add_page_break()
    doc.add_heading("四、需求内容详述", level=1)
    doc.add_heading("4.1 核心主流程", level=2)
    doc.add_picture(str(flow), width=Inches(6.65))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("图 2｜医院数据优先的核心闭环。", style="Caption")
    for text in [
        "用户在“数据接入”上传医院患者级 CSV，或选择内置合成数据进行演示。",
        "系统执行阻断性质量检查，生成 years_since_screen、Care Gap 标签/分数、完成筛查概率与检出概率，并展示字段血缘。",
        "用户主动点击“应用于评估与优先级仿真”，数据仅保留在当前 Streamlit 会话。",
        "AI 决策助手读取人群规模、平均年龄和当前筛查率；用户只补充目标筛查率。",
        "系统列出运行 ROI 所需参数的值、来源、说明和缺失项；缺失项按参数定向到医院数据、PubMed、SEER、CMS 或指南。",
        "用户保存候选值及来源，并勾选确认；系统调用 Python 版确定性 ROI 引擎（与原 R 公式一致）。",
        "如选择实时 AI，系统用固定 Prompt 生成结构化管理层报告；本地校验器验证数字、PMID、摘要摘录和 ROI 快照。",
        "通过校验后展示报告草稿及建议行动；失败则阻断，绝不展示未校验内容。",
    ]:
        add_numbered(doc, text)
    doc.add_heading("4.2 核心业务规则", level=2)
    add_table(doc, ["规则编号", "规则", "说明"], [
        ["BR-01", "医院数据优先", "不得在未读取医院数据时复用历史演示场景；数据变化后清空旧场景与旧报告。"],
        ["BR-02", "只补充缺失参数", "已由医院数据提供的参数不要求重复检索文献。"],
        ["BR-03", "人工确认后计算", "参数来源未确认时不得运行 ROI。"],
        ["BR-04", "确定性计算", "所有 ROI 数字由原 R 模型迁移的 Python 函数计算；LLM 不得计算或改写。"],
        ["BR-05", "证据不自动生效", "文献提取结果进入价值库不等于被 ROI 采用；须单独确认参数。"],
        ["BR-06", "报告先校验后展示", "任何无法追溯到 ROI 快照或批准证据的数字/PMID均阻断。"],
        ["BR-07", "明确模型边界", "所有核心页面展示“40–74岁女性DBT筛查”。"],
        ["BR-08", "不承诺真实效果", "报告使用“模拟/预计/情景输出”，不得写成已实现、保证或因果结论。"],
    ], [2.0, 4.0, 10.0], font_size=8.4)

    doc.add_heading("五、产品功能详细说明（核心）", level=1)
    doc.add_heading("5.1 模块一：AI 决策助手", level=2)
    doc.add_heading("5.1.1 页面入口与说明", level=3)
    add_kv_table(doc, [("入口路径", "/（首页）"), ("页面用途", "串联医院数据、缺失参数补充、人工确认、确定性 ROI 和 AI 报告的一站式主流程。")])
    doc.add_heading("5.1.2 页面结构与字段", level=3)
    add_table(doc, ["字段名", "类型", "规则", "默认值", "必填", "备注"], [
        ["目标筛查率", "数字输入（%）", "0–100；建议高于当前筛查率", "70%", "是", "当前筛查率由医院数据推导"],
        ["筛查方式", "只读文本", "V1.0 固定 DBT / 3D mammography", "DBT", "是", "不可切换"],
        ["癌症检出率", "数字输入（每千次）", "≥0；必须记录来源", "R 默认 6.2", "是", "医院缺失时补充"],
        ["召回率", "数字输入（%）", "0–100；必须记录来源", "R 默认 11.5%", "是", "医院缺失时补充"],
        ["来源类型", "下拉", "R默认/医院数据/PubMed/其他官方来源/情景假设", "R模型默认值", "是", "与参数绑定"],
        ["来源说明", "文本", "建议写 PMID、报表名、年份或假设理由", "", "否", "空时显示未补充说明"],
        ["参数确认", "复选框", "未勾选不得运行", "未选", "是", "人机协同门槛"],
        ["使用AI撰写报告", "复选框", "关闭时使用无 API 的验证演示报告", "开启", "否", "API失败不影响ROI结果"],
    ], [2.8, 2.5, 5.2, 2.5, 1.4, 3.0], font_size=7.6)
    doc.add_heading("5.1.3 交互与操作行为", level=3)
    add_bullet(doc, "“使用医院数据建立分析场景”：读取当前会话批准数据，生成参数计划；不上传数据时页面停止，不展示旧 Demo。")
    add_bullet(doc, "“保存该参数及来源”：将候选值、来源类型与来源说明存入会话状态，不立即运行计算。")
    add_bullet(doc, "“确认并运行 ROI”：校验确认勾选与输入范围，调用 calculate_breast_roi，保存只读快照并刷新结果卡。")
    add_bullet(doc, "“生成报告”：调用 AI 或演示生成器；成功后再经本地校验器验证，失败显示清晰错误。")
    doc.add_heading("5.1.4 AI Prompt 与校验约束", level=3)
    add_bullet(doc, "输入：医院场景、确定性 ROI 完整快照、人工批准证据、目标受众、输出语言。")
    add_bullet(doc, "输出结构：管理层摘要、临床影响、财务影响、证据解读、关键假设、局限性、建议行动、证据主张、PMID列表、ROI快照。")
    add_bullet(doc, "禁止事项：新增 ROI 数字、虚构 PMID、超出摘要证据下结论、将模拟结果描述为真实效果。")
    add_bullet(doc, "校验失败：自动重试一次（底层生成器）；仍失败则不展示未校验报告。")
    doc.add_heading("5.1.5 异常状态", level=3)
    add_table(doc, ["异常", "页面反馈", "兜底"], [
        ["无批准医院数据", "提示先进入数据接入", "停止页面流程，不加载历史场景"],
        ["缺失参数未保存", "显示待补充清单", "允许使用明确标注的R默认值/情景假设"],
        ["未确认来源", "提示先确认参数", "不运行ROI"],
        ["OpenAI Key缺失或API失败", "显示生成失败原因", "切换验证演示报告；ROI结果仍保留"],
        ["报告数字/引用不一致", "显示校验错误", "阻断展示并保留原ROI快照"],
    ], [4.0, 5.2, 7.0], font_size=8.4)

    doc.add_page_break()
    doc.add_heading("5.2 模块二：数据接入", level=2)
    doc.add_heading("5.2.1 页面入口与说明", level=3)
    add_kv_table(doc, [("入口路径", "/data_intake（导航：数据接入 → 上传与检查）"), ("页面用途", "接收医院基础字段，执行质量门禁并生成下游分析需要的派生字段。")])
    doc.add_heading("5.2.2 医院上传字段", level=3)
    add_table(doc, ["字段名", "类型", "规则", "默认值", "必填", "备注"], [
        ["patient_id", "文本", "每行非空且唯一", "", "是", "非合成数据下游使用前哈希"],
        ["as_of_date", "日期", "有效日期；不得早于最近筛查日期", "", "是", "数据基准日"],
        ["age", "整数", "40–74", "", "是", "超出模型范围阻断"],
        ["last_screen_date", "日期", "从未筛查者可空；其他人须有效", "", "条件必填", "用于计算距上次筛查年数"],
        ["never_screened", "布尔", "true / false 可解析", "", "是", "从未筛查标记"],
        ["has_active_appointment", "布尔", "true / false 可解析", "", "是", "有预约者不纳入当前缺口"],
        ["outreach_consent", "布尔", "true / false 可解析", "", "是", "不同意者不进入外展"],
        ["prior_abnormal", "布尔", "缺失时默认为 false 并告警", "false", "否", "排序规则特征"],
        ["family_history", "布尔", "缺失时默认为 false 并告警", "false", "否", "排序规则特征"],
        ["preferred_language", "文本", "可空", "", "否", "未来外展文案适配"],
        ["clinic_id", "文本", "可空", "", "否", "未来机构分组"],
    ], [2.8, 2.0, 5.1, 2.0, 1.8, 3.2], font_size=7.5)
    doc.add_heading("5.2.3 派生字段与血缘", level=3)
    add_table(doc, ["派生字段", "计算规则", "来源类型", "用途"], [
        ["years_since_screen", "(as_of_date − last_screen_date) / 365.25；从未筛查设为20年", "确定性规则", "缺口识别与排序"],
        ["ground_truth_gap", "从未筛查或距上次筛查≥筛查间隔；且无有效预约、同意外展", "演示真值规则", "演示评估/资源优化，不代表真实临床金标准"],
        ["care_gap_score", "0.15 + 0.55×逾期强度 + 0.12×从未筛查 + 0.10×既往异常 + 0.08×家族史 − 0.35×已有预约 − 0.30×不同意外展；截断至0–1", "透明业务规则", "外展优先级"],
        ["completion_probability", "基础完成率 +0.10×既往异常 +0.05×家族史 −0.08×从未筛查 −0.20×不同意外展；截断至0.05–0.90", "明确标注的演示假设", "估计完成筛查人数"],
        ["detection_probability", "R模型检出率/1000 × 年龄段发病率校正系数", "R模型参数", "估计检出病例"],
    ], [3.1, 8.7, 3.1, 2.4], font_size=7.5)
    doc.add_heading("5.2.4 交互与操作行为", level=3)
    add_bullet(doc, "支持上传 CSV 或选择内置 10,000 行合成测试文件；单文件上限 100,000 行。")
    add_bullet(doc, "质量门禁展示行数、列数、失败检查、警告、数据类型及逐项详情；存在失败项时禁止应用。")
    add_bullet(doc, "“应用于评估与优先级仿真”：保存至当前会话，清空旧的助手场景、结果和报告，防止数据串用。")
    add_bullet(doc, "“移除当前会话数据”：删除会话中的批准数据，不删除本地示例文件或证据库。")

    doc.add_heading("5.3 模块三：证据与参数", level=2)
    doc.add_heading("5.3.1 页面入口与说明", level=3)
    add_kv_table(doc, [("入口路径", "/evidence_value_library（导航：证据与参数）"), ("页面用途", "检索与模型参数直接相关的真实 PubMed 记录，生成结构化效应量候选，经人工审核后进入价值库。")])
    doc.add_heading("5.3.2 页面结构与字段", level=3)
    add_table(doc, ["字段名", "类型", "规则", "默认值", "必填", "备注"], [
        ["需要补充的证据", "下拉", "检出率/召回率/随访完成率/Stage shift 等", "癌症检出率", "是", "驱动检索模板"],
        ["PubMed检索式", "多行文本", "允许编辑；需与参数和DBT场景相关", "参数化模板", "是", "调用NCBI E-utilities"],
        ["文献选择", "下拉", "仅选择当前检索结果", "第一篇", "是", "展示标题、PMID、摘要"],
        ["审核说明", "多行文本", "批准/拒绝时必须说明适用范围或原因", "", "是", "形成审计记录"],
    ], [3.0, 2.2, 5.2, 2.7, 1.5, 3.0], font_size=7.6)
    doc.add_heading("5.3.3 交互与操作行为", level=3)
    add_bullet(doc, "“检索 PubMed”：返回真实 PubMed 记录；请求失败时显示错误，不生成虚构文献。")
    add_bullet(doc, "“AI 提取效应量”：仅根据所选摘要生成结构化候选，提取研究设计、人群、结局、效应值、适用参数与局限性。")
    add_bullet(doc, "“批准并存入价值库 / 拒绝候选结果”：必须填写审核说明；AI 候选不会自动修改 ROI。")
    add_bullet(doc, "价值库仅保存经过审核的记录，并保留审核状态与审计日志。")
    add_callout(doc, "来源边界", "PubMed 适合临床研究与系统综述；SEER 更适合发病率和分期分布，CMS/医院财务更适合成本，USPSTF 等指南更适合筛查年龄与频率。V1.0 只自动接入 PubMed，其余来源提供定向链接和人工录入。", "info")

    doc.add_heading("5.4 模块四：高级 ROI 仿真", level=2)
    doc.add_heading("5.4.1 页面入口与说明", level=3)
    add_kv_table(doc, [("入口路径", "/roi_simulation（导航：决策分析 → 高级 ROI 仿真）"), ("页面用途", "供分析师复核参数、开展 what-if 和敏感性分析；普通用户优先使用 AI 决策助手。")])
    doc.add_heading("5.4.2 核心输入与公式", level=3)
    add_table(doc, ["类别", "字段", "默认/来源", "规则"], [
        ["普通区域", "人群规模、当前筛查率、目标筛查率", "优先读取AI助手确认场景；否则R默认值", "改变后实时计算"],
        ["高级假设", "年龄、筛查成本/间隔、召回/随访、检出率、生命挽救率", "原R默认或人工确认值", "默认折叠"],
        ["高级假设", "分期占比、分期成本、阶段迁移比例", "原R默认或人工确认值", "默认折叠；用于敏感性分析"],
        ["输出", "新增筛查、检出病例、生命挽救、净节约、ROI", "确定性函数", "不得由LLM修改"],
    ], [2.3, 5.3, 5.1, 4.0], font_size=8.0)
    add_callout(doc, "主要公式", "新增筛查 = 人群 × max(目标率−当前率,0)；检出病例 = 新增筛查 × 每千次检出率/1000 × 年龄校正；净节约 = 避免治疗成本 − 筛查成本 − 随访成本；ROI = 净节约 / 项目成本。", "success")
    doc.add_heading("5.4.3 交互与操作行为", level=3)
    add_bullet(doc, "默认读取 AI 决策助手已确认的快照；若无快照，明确提示使用 R 默认值作为演示起点。")
    add_bullet(doc, "高级参数放入折叠区，避免普通用户被医学细节干扰。")
    add_bullet(doc, "参数变化立即更新 KPI、明细、图表和敏感性分析；页面始终标注模型范围与确定性计算属性。")

    doc.add_page_break()
    doc.add_heading("5.5 模块五：外展资源优化", level=2)
    doc.add_heading("5.5.1 页面入口与说明", level=3)
    add_kv_table(doc, [("入口路径", "/risk_prioritization（导航：决策分析 → 外展资源优化）"), ("页面用途", "在联系名额有限时，对比随机联系与透明规则优先联系，帮助运营决定先联系哪些人。")])
    doc.add_heading("5.5.2 页面结构与字段", level=3)
    add_table(doc, ["字段名", "类型", "规则", "默认值", "必填", "备注"], [
        ["本轮最多联系人数", "数字", "100–当前可分析人群", "min(2000, 人群数)", "是", "业务参数"],
        ["每联系一人成本", "金额", "0–500美元", "$12", "是", "短信/电话/人工成本"],
        ["随机方案重复次数", "滑块", "20–500", "100", "否", "技术设置，默认折叠"],
        ["可重复实验编号", "整数", "0–100,000", "42", "否", "保证演示可复现"],
    ], [3.2, 2.2, 4.2, 3.0, 1.6, 3.7], font_size=7.8)
    doc.add_heading("5.5.3 排序与价值计算", level=3)
    add_callout(doc, "优先级分数", "60% × Care Gap 分数 + 25% × 距上次筛查时间标准化分数 + 15% × 预计完成筛查概率。系统对全部人群计算分数，再选取得分最高的 N 人作为本轮建议名单。", "info")
    add_bullet(doc, "系统不判断谁一定会患癌，不自动联系患者，也不替代医院的最终运营决策。")
    add_bullet(doc, "随机基准重复抽样并取均值；优先策略为确定性排序。")
    add_bullet(doc, "预期完成筛查和预期检出病例均为概率加总的期望值，不是实际观察人数。")
    add_bullet(doc, "经济结果复用当前 ROI 引擎中的年度筛查成本、随访成本与每例阶段迁移节约。")

    doc.add_heading("5.6 跨模块状态与数据流", level=2)
    add_table(doc, ["上游", "会话状态 / 数据", "下游", "刷新规则"], [
        ["数据接入", "approved_population + metadata", "AI助手、外展资源优化", "新数据应用后清空旧场景/报告"],
        ["AI助手", "roi_inputs_snapshot + roi_results_snapshot", "高级ROI仿真、报告", "参数确认并运行后更新"],
        ["证据与参数", "Care Gap Value Library", "AI报告、参数来源复核", "仅Approved记录可用"],
        ["高级ROI仿真", "当前ROI输入/结果", "图表、敏感性分析", "组件变化即时更新"],
    ], [3.0, 5.0, 4.2, 4.2], font_size=8.2)

    doc.add_heading("六、原型与交互说明", level=1)
    doc.add_heading("6.1 原型链接", level=2)
    add_kv_table(doc, [("Axure / Figma", ""), ("本地演示地址", "http://localhost:8502/（仅本机运行时有效）"), ("代码仓库", "")])
    doc.add_heading("6.2 全局交互原则", level=2)
    add_bullet(doc, "默认中文，保留 English 切换；内部字段名和模型键值不翻译，避免计算链路变化。")
    add_bullet(doc, "主导航以任务而非技术命名；专业能力放入二级入口或折叠区。")
    add_bullet(doc, "所有关键状态使用明确反馈：绿色=可继续，黄色=假设/提醒，红色=阻断或校验失败。")
    add_bullet(doc, "无数据、无证据、无 API Key、验证失败均提供下一步操作，不展示空白页或虚构内容。")
    add_bullet(doc, "所有概率、模拟值和合成数据必须明确标注，不与真实医院效果混淆。")
    doc.add_heading("6.3 关键页面跳转", level=2)
    add_table(doc, ["当前页面", "触发条件", "目标页面/区域", "上下文保留"], [
        ["AI决策助手", "无批准数据", "数据接入", "不保留旧场景"],
        ["数据接入", "应用成功", "返回AI决策助手", "批准数据保留在会话"],
        ["AI决策助手", "缺失检出率/召回率", "PubMed定向检索或证据与参数", "参数名称与建议来源可见"],
        ["AI决策助手", "确认并运行ROI", "本页ROI结果", "保存输入与结果快照"],
        ["AI决策助手", "生成报告", "本页管理层报告草稿", "绑定当前ROI快照和批准证据"],
    ], [3.4, 4.0, 4.8, 4.2], font_size=8.2)

    doc.add_heading("七、非功能需求（NFR）", level=1)
    doc.add_heading("7.1 性能", level=2)
    add_table(doc, ["项目", "目标 / 约束", "状态"], [
        ["确定性ROI计算", "单次本地计算目标 <1秒", "待正式压测；当前交互为即时刷新"],
        ["CSV处理", "≤100,000行；质量检查期间显示进度/状态", "已设置行数上限"],
        ["PubMed检索", "外部请求建议超时≤30秒，失败可重试并显示错误", "当前依赖网络"],
        ["AI报告生成", "建议端到端≤45秒；展示生成中状态", "待建立延迟埋点"],
        ["并发支持", "", "待部署方案与压测后确定"],
    ], [4.0, 8.5, 4.0], font_size=8.5)
    doc.add_heading("7.2 兼容性", level=2)
    add_kv_table(doc, [("APP版本", "Web MVP，无原生 APP"), ("浏览器", "Chrome 最新稳定版；Safari 最新稳定版待完整回归"), ("机型", "优先桌面端 1440×900 及以上；移动端仅基础浏览，非本期重点")])
    doc.add_heading("7.3 安全与风控", level=2)
    add_bullet(doc, "权限控制：MVP 暂无账号体系；生产化前必须增加角色权限、机构隔离和审计访问控制。")
    add_bullet(doc, "敏感信息处理：上传文件仅保留在当前会话；非合成数据患者 ID 在下游使用前执行本地单向哈希；不写入数据库。")
    add_bullet(doc, "API Key：仅存放于 .streamlit/secrets.toml 或部署环境密钥，严禁提交 GitHub 或展示在页面日志。")
    add_bullet(doc, "防重复提交：生成类按钮执行期间应禁用；同一输入可使用会话缓存，外部接口失败可有限重试。")
    add_bullet(doc, "内容安全：不提供诊断或医疗建议；报告必须包含模型性质、证据边界和局限性。")
    doc.add_heading("7.4 可用性与稳定性", level=2)
    add_bullet(doc, "降级策略：OpenAI 不可用时切换验证演示报告；PubMed 不可用时显示错误并保留手工来源入口；ROI 与数据质检仍可使用。")
    add_bullet(doc, "容灾策略：MVP 为本地/单实例会话应用，暂无生产容灾；正式部署前需配置健康检查、日志、密钥轮换和数据备份策略。")
    add_bullet(doc, "可恢复性：外部 API 失败不得清空已确认参数与已计算 ROI 快照。")

    doc.add_heading("八、数据埋点与分析", level=1)
    add_callout(doc, "实现状态", "以下为 P1 运营分析需求，当前 MVP 尚未接入正式埋点平台。任何埋点不得上传患者级字段、原始文本或可识别信息。", "warning")
    doc.add_heading("8.1 埋点事件清单", level=2)
    add_table(doc, ["事件名", "触发时机", "核心属性", "备注"], [
        ["data_upload_completed", "CSV读取完成", "row_count, schema_type, synthetic_flag", "不记录文件内容/患者ID"],
        ["data_quality_completed", "质检结束", "fail_count, warning_count, duration_ms", "衡量数据可用性"],
        ["dataset_applied", "用户应用数据集", "row_count, deidentified_flag", "核心漏斗起点"],
        ["scenario_created", "助手建立场景", "known_param_count, missing_param_count", "衡量医院数据完整度"],
        ["evidence_search_completed", "PubMed检索完成", "topic, result_count, latency_ms, success", "不记录完整查询中的敏感信息"],
        ["parameter_confirmed", "确认参数并运行ROI", "source_type, parameter_count", "不记录患者级值"],
        ["roi_run_completed", "ROI返回结果", "success, duration_ms, model_version", "金额可按区间上报"],
        ["report_generation_completed", "报告成功或被阻断", "mode, latency_ms, validation_status, error_type", "记录token/cost需去敏"],
        ["outreach_simulation_completed", "资源优化完成", "population_bucket, capacity_ratio", "不记录联系人ID"],
    ], [4.1, 4.0, 6.5, 3.1], font_size=7.5)
    doc.add_heading("8.2 指标口径定义", level=2)
    add_bullet(doc, "核心流程完成率 = 成功生成经校验报告的会话数 ÷ 已应用数据集的会话数。")
    add_bullet(doc, "数据一次通过率 = 首次上传即通过阻断性质量检查的数据集数 ÷ 完成上传的数据集数。")
    add_bullet(doc, "参数缺失率 = 场景缺失参数数 ÷ 运行 ROI 所需参数总数。")
    add_bullet(doc, "AI 报告校验通过率 = 首次通过数字与引用校验的报告数 ÷ AI 返回报告数。")
    add_bullet(doc, "外部接口成功率 = 成功请求数 ÷ 请求总数，分别统计 PubMed 与 OpenAI。")
    add_bullet(doc, "平均报告成本 = 报告生成总 API 成本 ÷ 成功报告数（待接入 token 与价格配置）。")
    doc.add_heading("8.3 数据看板 / 报表需求", level=2)
    add_table(doc, ["看板", "指标", "受众", "优先级"], [
        ["产品漏斗", "上传→应用→场景→ROI→报告", "产品经理", "P1"],
        ["数据质量", "失败项分布、一次通过率、字段缺失率", "数据分析师", "P1"],
        ["AI质量", "校验通过率、失败类型、重试率、PMID违规率", "AI产品经理", "P1"],
        ["性能与成本", "P50/P95延迟、token、单报告成本、接口错误率", "产品/技术", "P1"],
    ], [4.0, 7.0, 3.0, 2.3], font_size=8.4)

    doc.add_heading("九、上线与运营", level=1)
    doc.add_heading("9.1 灰度策略", level=2)
    add_bullet(doc, "阶段1：内部演示，仅使用内置合成数据和验证演示报告。")
    add_bullet(doc, "阶段2：邀请制测试，允许上传已脱敏/合成 CSV，开启 PubMed 与 OpenAI 实时能力。")
    add_bullet(doc, "阶段3：在完成权限、隐私、监控和安全评审后，评估真实医院试点；不得直接沿用 MVP 的会话级数据方案。")
    doc.add_heading("9.2 上线检查清单", level=2)
    add_table(doc, ["类别", "检查项", "通过标准"], [
        ["配置", "OPENAI_API_KEY / OPENAI_MODEL / 网络访问", "密钥可用且不出现在代码与日志"],
        ["数据", "CSV模板、合成标识、年龄/日期规则", "标准测试集通过；真实数据路径完成去标识化验证"],
        ["模型", "R/Python回归、默认场景与边界场景", "关键公式结果一致，测试全部通过"],
        ["AI", "数字、PMID、摘要摘录、建议行动", "任何失败结果均被阻断"],
        ["交互", "无数据、API失败、报告失败、参数未确认", "均有中文可操作提示"],
        ["监控", "错误、延迟、外部接口、成本", "P1上线前接入基础日志与告警"],
    ], [3.0, 7.2, 6.3], font_size=8.3)
    doc.add_heading("9.3 运营配置", level=2)
    add_kv_table(doc, [
        ("文案", "模型范围、合成数据、情景估算、非医疗建议、证据不足提示可配置但不可删除核心风险含义。"),
        ("规则", "筛查间隔、基础完成率假设、外展名额和成本可配置；权重规则修改需版本化并回归。"),
        ("开关", "实时 AI / 验证演示、语言、技术设置折叠、证据专业入口。"),
    ])
    doc.add_heading("9.4 客服 FAQ", level=2)
    add_table(doc, ["问题", "标准回答"], [
        ["为什么必须先上传数据？", "系统采用医院数据优先原则，避免用演示默认值替代本地人群与筛查现状。"],
        ["一定要找 PubMed 吗？", "不一定。医院已提供的数据直接使用；PubMed只用于补充适合临床文献的缺失参数。成本、发病率和指南应去对应官方来源。"],
        ["AI会修改ROI吗？", "不会。AI只负责非结构化信息提取和报告表达，ROI由确定性公式计算。"],
        ["为什么报告生成失败？", "可能是API、余额、模型配置或数字/引用校验失败。系统不会显示未校验报告，可切换验证演示。"],
        ["外展名单是患癌预测吗？", "不是。它是筛查缺口与完成可能性的运营优先级建议。"],
        ["结果能代表医院会节省这些钱吗？", "不能。结果是情景模拟，需用真实参与率、随访、成本和长期结果验证。"],
    ], [5.0, 11.5], font_size=8.4)

    doc.add_heading("十、风险与兜底方案", level=1)
    doc.add_heading("10.1 技术风险", level=2)
    add_table(doc, ["风险", "影响", "预防", "兜底"], [
        ["外部API不可用", "无法检索或生成AI报告", "超时、错误分类、有限重试", "保留ROI；切换演示报告/人工来源"],
        ["Streamlit会话丢失", "上传数据与快照消失", "明确会话级存储说明", "重新上传；生产版引入安全持久化"],
        ["R/Python公式漂移", "结果不可复现", "回归测试与模型版本号", "回滚上一版本计算模块"],
        ["大文件或并发", "页面卡顿", "100k行限制与缓存", "拒绝超限文件；排队/扩容待生产设计"],
    ], [4.0, 4.2, 5.2, 4.1], font_size=7.9)
    doc.add_heading("10.2 业务风险", level=2)
    add_bullet(doc, "参数适用性错误：同一研究可能与本地年龄、设备、流程不同。兜底为显示人群/研究设计/局限性并要求人工确认。")
    add_bullet(doc, "结果被过度解读：强制使用“模拟/预计”，报告固定包含关键假设、局限性和建议验证项。")
    add_bullet(doc, "优先级造成不公平：P1 增加按年龄、语言、诊所等分组的覆盖率与误差审查，不允许系统自动外展。")
    doc.add_heading("10.3 合规风险", level=2)
    add_bullet(doc, "患者隐私：MVP 不应接收未获授权的真实 PHI；生产试点前完成法务、隐私、安全与数据处理协议评审。")
    add_bullet(doc, "医疗器械/临床决策边界：产品定位为运营与价值情景决策支持，不用于诊断或替代医生判断。")
    add_bullet(doc, "文献版权与引用：仅展示检索元数据与合规摘要内容，保留 PMID 和原文链接，不虚构引用。")
    doc.add_heading("10.4 回滚 / 降级方案", level=2)
    add_bullet(doc, "AI报告异常：关闭实时 AI 开关，回退验证演示模板；ROI引擎继续可用。")
    add_bullet(doc, "证据模块异常：隐藏证据入口，允许人工填写来源；不得自动赋值。")
    add_bullet(doc, "新模型版本异常：回滚至上一版本计算模块及参数文件，保留输入/输出快照用于审计。")
    add_bullet(doc, "发现隐私风险：立即停止上传入口并清除会话数据；生产环境按事件响应流程处理。")

    doc.add_heading("十一、附录", level=1)
    doc.add_heading("11.1 名词解释", level=2)
    add_table(doc, ["术语", "解释"], [
        ["DBT", "Digital Breast Tomosynthesis，数字乳腺断层合成/三维乳腺X线摄影。"],
        ["Care Gap", "应接受筛查但尚未完成的服务缺口；本MVP使用透明规则定义。"],
        ["ROI", "净节约 ÷ 筛查项目成本；为模型情景指标，不代表实际已实现收益。"],
        ["召回率", "筛查后因异常或需进一步确认而被要求追加检查的比例，不是营销召回。"],
        ["检出率", "每1,000次筛查发现的乳腺癌病例数。"],
        ["Stage shift", "假设部分病例由较晚分期转向较早分期；当前证据不足时仅作敏感性假设。"],
        ["PMID", "PubMed给每篇收录文献分配的唯一编号。"],
        ["RAG", "检索增强生成；先获取限定证据，再让大模型基于证据回答。当前独立问答入口已隐藏，底层能力保留。"],
        ["确定性模型", "同样输入得到同样输出的公式模型，与大模型文本生成不同。"],
        ["合成数据", "为演示构造、不对应真实患者的数据。"],
    ], [3.4, 13.0], font_size=8.5)
    doc.add_heading("11.2 参考文档", level=2)
    add_bullet(doc, "项目需求：APEX Project Brief_3_25_26.pdf")
    add_bullet(doc, "原始模型：breast_with_baseline_delta.R")
    add_bullet(doc, "当前工程：breast_roi_copilot/README.md")
    add_bullet(doc, "计算模块：src/models/breast_roi.py；输入输出：src/models/schemas.py")
    add_bullet(doc, "数据规则：src/data_intake/validator.py；优先级规则：src/prioritization/simulator.py")
    add_bullet(doc, "报告 Prompt 与校验：src/reporting/prompts.py、src/reporting/validator.py")
    add_bullet(doc, "竞品分析：")
    add_bullet(doc, "需求池：")
    add_bullet(doc, "接口文档：")
    doc.add_heading("11.3 常见问题 Q&A", level=2)
    add_table(doc, ["问题", "回答"], [
        ["这是 workflow 还是 Agent？", "V1.0 是受控 AI workflow：步骤固定、关键节点人工确认。它可描述为有限自主的决策助手，但不应包装为可自由规划和执行的全自主 Agent。"],
        ["哪里用了 RAG？", "证据检索与基于已检索/已批准摘要的分析采用检索增强思路；独立循证问答页面已隐藏，但检索、结构化提取和引用校验能力保留。"],
        ["哪里用了 Prompt？", "管理层报告和证据提取使用固定系统提示词、结构化输出约束和禁止事项；自然语言与Codex沟通是开发过程，不等同于产品内Prompt版本实验。"],
        ["哪里体现 AI 产品经理能力？", "体现在能力边界、医院数据优先流程、Prompt结构化输出、人机确认、引用/数字校验、失败降级、指标设计和风险治理。"],
        ["为什么不让 AI 算 ROI？", "ROI是可审计公式，确定性引擎更稳定、可测试；LLM更适合处理文献和报告表达。"],
        ["模型性能评估页面为什么隐藏？", "其评估的是合成Care Gap规则分数，不是LLM质量，容易干扰主线；代码保留供专业复核。"],
    ], [5.0, 11.5], font_size=8.3)
    doc.add_heading("11.4 验收用例摘要", level=2)
    add_table(doc, ["用例", "输入/操作", "预期结果"], [
        ["无数据进入助手", "清空会话后打开首页", "提示先上传，不显示历史10万人场景"],
        ["医院CSV通过", "上传标准10,000行合成文件", "0个阻断失败；可应用；显示派生字段血缘"],
        ["医院CSV失败", "缺少patient_id或年龄越界", "显示失败项；应用按钮不可用"],
        ["参数未确认", "不勾选确认直接运行", "阻断并提示确认来源"],
        ["ROI回归", "使用R默认输入", "关键输出与R基线一致"],
        ["报告含非法数字", "AI返回快照外数字", "本地校验拦截，不显示报告"],
        ["报告含非法PMID", "AI返回不在批准证据中的PMID", "本地校验拦截"],
        ["API不可用", "关闭网络或使用无效Key", "ROI保留；显示错误；可切换验证演示"],
        ["外展资源优化", "10,000人、2,000名额", "对全体评分，仅输出最高2,000人为建议名单并与随机基准对比"],
    ], [4.2, 6.0, 6.3], font_size=8.0)

    doc.add_paragraph()
    add_callout(doc, "文档结束", "本 PRD 为 V1.0 初稿。撰写人、原型链接、代码仓库、生产并发与真实上线数据待项目负责人补充。", "success")
    apply_cjk_font(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
